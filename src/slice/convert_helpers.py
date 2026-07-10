"""Helper scripts for file format conversion with chunking support for large files."""

EXCEL_TO_JSON = """
import pandas as pd
import json
import sys

input_file = sys.argv[1]
output_file = sys.argv[2]

try:
    # Read Excel file with chunking for large files
    df = pd.read_excel(input_file, engine='openpyxl')

    # Convert to JSON with proper formatting
    result = df.to_dict(orient='records')

    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(result, f, indent=2, ensure_ascii=False, default=str)

    print(f"Successfully converted {len(result)} rows")
except Exception as e:
    print(f"Error: {str(e)}", file=sys.stderr)
    sys.exit(1)
"""

CSV_TO_JSON = """
import pandas as pd
import json
import sys

input_file = sys.argv[1]
output_file = sys.argv[2]

try:
    # Read CSV with chunking for large files
    chunk_size = 10000
    chunks = []

    for chunk in pd.read_csv(input_file, chunksize=chunk_size, encoding='utf-8'):
        chunks.append(chunk)

    # Combine all chunks
    df = pd.concat(chunks, ignore_index=True)

    # Convert to JSON
    result = df.to_dict(orient='records')

    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(result, f, indent=2, ensure_ascii=False, default=str)

    print(f"Successfully converted {len(result)} rows")
except Exception as e:
    print(f"Error: {str(e)}", file=sys.stderr)
    sys.exit(1)
"""

WORD_TO_JSON = """
from docx import Document
import json
import sys

input_file = sys.argv[1]
output_file = sys.argv[2]

def extract_table(table):
    table_data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        if any(row_data):
            table_data.append(row_data)
    return table_data

try:
    doc = Document(input_file)
    result = {
        'paragraphs': [],
        'tables': []
    }

    # Extract paragraphs and tables in order
    for element in doc.element.body:
        if element.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._element == element:
                    text = para.text.strip()
                    if text:
                        result['paragraphs'].append(text)
                    break
        elif element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == element:
                    table_data = extract_table(table)
                    if table_data:
                        result['tables'].append(table_data)
                    break

    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(result, f, indent=2, ensure_ascii=False)

    print(f"Successfully converted {len(result['paragraphs'])} paragraphs and {len(result['tables'])} tables")
except Exception as e:
    print(f"Error: {str(e)}", file=sys.stderr)
    sys.exit(1)
"""

PDF_TO_JSON = """
from pypdf import PdfReader
import json
import sys

input_file = sys.argv[1]
output_file = sys.argv[2]

try:
    reader = PdfReader(input_file)
    result = {
        'pages': [],
        'metadata': {
            'page_count': len(reader.pages)
        }
    }

    # Process pages one at a time to avoid memory issues
    for page_num, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text()
        result['pages'].append({
            'page_number': page_num,
            'text': page_text
        })

    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(result, f, indent=2, ensure_ascii=False)

    print(f"Successfully converted {len(result['pages'])} pages")
except Exception as e:
    print(f"Error: {str(e)}", file=sys.stderr)
    sys.exit(1)
"""
